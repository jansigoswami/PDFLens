import streamlit as st
import os
import time
import tempfile
from openai import RateLimitError
from dotenv import load_dotenv
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import PyPDF2

# Load environment variables from .env file
load_dotenv()

from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.prompts import PromptTemplate
from langchain_cerebras import ChatCerebras
from langchain.schema import Document as LangChainDocument

# ---------------------- Streamlit UI ----------------------
st.set_page_config(page_title="File-Based Chatbot", layout="wide")
st.title("📘 Smart Document Chatbot (with Page Reference)")
st.markdown("Ask questions based on the uploaded PDF, Word, or Excel file.")

# ---------------------- File Upload ----------------------
uploaded_file = st.file_uploader("Upload your document (PDF, Word, or Text)", type=["pdf", "docx", "txt"])

# ---------------------- PDF Loader with Correct Page Numbers ----------------------
def load_pdf_document(file_path):
    """Load PDF and extract text with correct 1-indexed page numbers"""
    pdf_reader = PdfReader(file_path)
    documents = []
    
    # First, extract all pages with their raw content
    raw_pages = []
    for i, page in enumerate(pdf_reader.pages, 1):
        try:
            text = page.extract_text()
            if text.strip():
                raw_pages.append({
                    'page_num': i,  # 1-based page numbering
                    'content': text,
                    'source': os.path.basename(file_path)
                })
        except Exception as e:
            st.warning(f"Warning: Could not extract text from page {i}: {str(e)}")
    
    # Now process each page and create documents
    for page_data in raw_pages:
        # Create a document for each page with accurate page number
        documents.append(
            LangChainDocument(
                page_content=page_data['content'],
                metadata={
                    'page': page_data['page_num'],  # Use the actual PDF page number
                    'source': page_data['source'],
                    'total_pages': len(pdf_reader.pages),
                    'is_pdf': True  # Add a flag to identify PDF documents
                }
            )
        )
    
    return documents

# ---------------------- Text File Loader ----------------------
def load_text_file(file_path):
    """Load text file and split into pages"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split content into chunks (pages) of approximately 1000 characters
        chunk_size = 1000
        chunks = [content[i:i + chunk_size] for i in range(0, len(content), chunk_size)]
        
        documents = []
        for i, chunk in enumerate(chunks, 1):
            if chunk.strip():  # Only add non-empty chunks
                documents.append(
                    LangChainDocument(
                        page_content=chunk,
                        metadata={
                            'page': i,
                            'source': os.path.basename(file_path),
                            'total_pages': len(chunks)
                        }
                    )
                )
        
        return documents
    except Exception as e:
        st.error(f"Error loading text file: {e}")
        return []

# ---------------------- Word Document Loader ----------------------
def load_word_document(file_path):
    """Load Word document and split by actual page breaks"""
    doc = DocxDocument(file_path)
    
    pages = []
    current_page_content = []
    current_page_num = 1
    
    for para in doc.paragraphs:
        # Check if this paragraph has a page break after it
        has_page_break = False
        
        for run in para.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                has_page_break = True
                break
        
        if para.text.strip():
            current_page_content.append(para.text.strip())
        
        if has_page_break:
            if current_page_content:
                pages.append({
                    'page': current_page_num,
                    'content': '\n'.join(current_page_content)
                })
                current_page_content = []
                current_page_num += 1
    
    # Add last page
    if current_page_content:
        pages.append({
            'page': current_page_num,
            'content': '\n'.join(current_page_content)
        })
    
    if not pages:
        all_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
        if all_text:
            pages.append({'page': 1, 'content': all_text})
    
    documents = []
    for page_data in pages:
        documents.append(
            LangChainDocument(
                page_content=page_data['content'],
                metadata={'page': page_data['page'], 'source': file_path}
            )
        )
    
    return documents

# ---------------------- Smart Chunker with Page Preservation ----------------------
def smart_chunk_with_pages(documents, chunk_size=1000, chunk_overlap=200):
    """Split documents into chunks while preserving exact page numbers"""
    chunks = []
    
    for doc in documents:
        page_num = doc.metadata.get('page', 1)
        content = doc.page_content
        
        if len(content) <= chunk_size:
            chunks.append(
                LangChainDocument(
                    page_content=content,
                    metadata={'page': page_num, 'source': doc.metadata.get('source', '')}
                )
            )
        else:
            # Split long content but keep same page number
            words = content.split()
            current_chunk = []
            
            for word in words:
                current_chunk.append(word)
                chunk_text = ' '.join(current_chunk)
                
                if len(chunk_text) >= chunk_size:
                    chunks.append(
                        LangChainDocument(
                            page_content=chunk_text,
                            metadata={'page': page_num, 'source': doc.metadata.get('source', '')}
                        )
                    )
                    # Keep overlap
                    overlap_size = max(1, int(len(current_chunk) * (chunk_overlap / chunk_size)))
                    current_chunk = current_chunk[-overlap_size:]
            
            if current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append(
                    LangChainDocument(
                        page_content=chunk_text,
                        metadata={'page': page_num, 'source': doc.metadata.get('source', '')}
                    )
                )
    
    return chunks

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
        tmp_file.write(uploaded_file.read())
        temp_path = tmp_file.name

    file_extension = uploaded_file.name.split(".")[-1].lower()
    
    try:
        if file_extension == "pdf":
            documents = load_pdf_document(temp_path)
            st.success(f"✅ Loaded PDF with {len(documents)} page(s)")
            
            with st.expander("📄 PDF Page Structure"):
                for doc in documents[:5]:  # Show first 5 pages
                    page_num = doc.metadata.get('page')
                    word_count = len(doc.page_content.split())
                    preview = doc.page_content[:100].replace('\n', ' ') + "..." if len(doc.page_content) > 100 else doc.page_content
                    st.write(f"Page {page_num}: {word_count} words")
                    st.caption(preview)
                    st.divider()
            
            chunks = smart_chunk_with_pages(documents, chunk_size=1000, chunk_overlap=200)
            
        elif file_extension == "docx":
            documents = load_word_document(temp_path)
            st.success(f"✅ Loaded Word document with {len(documents)} page(s)")
            
            with st.expander("📄 Document Page Structure"):
                for doc in documents:
                    page_num = doc.metadata.get('page')
                    word_count = len(doc.page_content.split())
                    preview = doc.page_content[:150] + "..."
                    st.write(f"Page {page_num}: {word_count} words")
                    st.caption(preview)
                    st.divider()
            
            chunks = smart_chunk_with_pages(documents, chunk_size=1000, chunk_overlap=200)
            
        elif file_extension == "txt":
            documents = load_text_file(temp_path)
            st.success(f"✅ Loaded text file with {len(documents)} sections")
            
            with st.expander("📄 Text File Content"):
                for doc in documents[:5]:  # Show first 5 sections
                    section_num = doc.metadata.get('page')
                    word_count = len(doc.page_content.split())
                    preview = doc.page_content[:150].replace('\n', ' ') + "..."
                    st.write(f"Section {section_num}: {word_count} words")
                    st.caption(preview)
                    st.divider()
            
            chunks = smart_chunk_with_pages(documents, chunk_size=1000, chunk_overlap=200)
            
        else:
            st.error("Unsupported file type. Please upload PDF or DOCX.")
            st.stop()

        st.info(f"📊 Created {len(chunks)} chunks for processing")

        # Verify chunks have correct page numbers
        with st.expander("🔍 Chunk Distribution by Page"):
            from collections import Counter
            page_counts = Counter([chunk.metadata.get('page') for chunk in chunks])
            for page, count in sorted(page_counts.items()):
                st.write(f"Page {page}: {count} chunk(s)")

        # ---------------------- Embeddings & Vectorstore ----------------------
        with st.spinner("Creating embeddings..."):
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            vectorstore = FAISS.from_documents(chunks, embeddings)
            retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

        # ---------------------- Cerebras LLM Setup ----------------------
        CEREBRAS_API_KEY = os.getenv("CEREBRAS_API_KEY")
        if not CEREBRAS_API_KEY:
            st.error("Cerebras API key not found. Please check your .env file.")
            st.stop()

        llm = ChatCerebras(
            api_key=CEREBRAS_API_KEY,
            model="llama3.1-8b",
            temperature=0,
            max_tokens=800
        )

        # ---------------------- Prompt Template ----------------------
        prompt = PromptTemplate(
            input_variables=["context", "question"],
            template=(
                "You are an assistant that answers based on a company document.\n"
                "Use only the context given below to answer accurately.\n\n"
                "Context:\n{context}\n\nQuestion: {question}\nAnswer:"
            ),
        )

        # ---------------------- Helper Functions ----------------------
        def get_source_pages(docs):
            """Extract unique page numbers from retrieved documents"""
            pages = set()
            for doc in docs:
                page_num = doc.metadata.get("page")
                if page_num is not None:
                    try:
                        pages.add(int(page_num))
                    except (ValueError, TypeError):
                        pass
            return sorted(list(pages)) if pages else [1]

        def get_most_relevant_page(docs):
            """Get the page from the highest-scoring retrieved document"""
            if not docs:
                return 1
            
            page_num = docs[0].metadata.get("page", 1)
            
            try:
                return int(page_num)
            except (ValueError, TypeError):
                return 1

        # ---------------------- Main Function ----------------------
        def ask_question(query):
            try:
                docs = retriever.invoke(query)
                
                if not docs:
                    return "No relevant information found in the document.", 1, []
                
                context = "\n\n".join([d.page_content for d in docs])
                formatted_prompt = prompt.format(context=context, question=query)

                response = None
                for attempt in range(5):
                    try:
                        response = llm.invoke(formatted_prompt)
                        break
                    except RateLimitError:
                        wait_time = (attempt + 1) * 5
                        st.warning(f"Server busy, retrying in {wait_time}s...")
                        time.sleep(wait_time)
                
                if response is None:
                    return "Rate limit reached. Try again later.", 1, []

                answer = response.content.strip()
                primary_page = get_most_relevant_page(docs)
                all_pages = get_source_pages(docs)
                
                return answer, primary_page, all_pages

            except Exception as e:
                st.error(f"Error in ask_question: {e}")
                return f"Error: {e}", 1, []

        # ---------------------- Chat Interface ----------------------
        st.divider()
        query = st.text_input("💬 Ask a question about the uploaded file:")

        if query:
            with st.spinner("🔍 Searching for the answer..."):
                answer, primary_page, all_pages = ask_question(query)

            st.markdown("### 🧠 Answer")
            st.write(answer)
            
            st.markdown("### 📄 Source Reference")
            st.info(f"Primary Page: {primary_page}")
            
            if len(all_pages) > 1:
                st.caption(f"Additional relevant pages: {', '.join(map(str, all_pages))}")
            
            with st.expander("🔍 View source context"):
                docs = retriever.invoke(query)
                for i, doc in enumerate(docs[:3]):
                    page = doc.metadata.get("page", "Unknown")
                    st.markdown(f"Chunk {i+1} (Page {page}):")
                    st.text(doc.page_content[:400] + "..." if len(doc.page_content) > 400 else doc.page_content)
                    if i < 2:
                        st.divider()

    except Exception as e:
        st.error(f"❌ Error: {e}")
        import traceback
        st.code(traceback.format_exc())
    
    finally:
        try:
            os.unlink(temp_path)
        except:
            pass